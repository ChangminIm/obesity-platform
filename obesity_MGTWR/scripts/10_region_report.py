# -*- coding: utf-8 -*-
"""후보지 30곳 × MGTWR 교차 검토 → 워드 보고서 + 엑셀 산출
   입력: docx 후보지 리스트, local_coefs_ob_MGTWR.csv, enp_critical_t_ob.csv
   출력: output/지역선정_MGTWR교차검토_260729.docx / .xlsx"""
import os
import pandas as pd
import docx as docxlib

BASE = r"C:\Users\User\vibe_coding\obesity_platform\obesity_MGTWR"
OUT = os.path.join(BASE, "output")
SRC_DOCX = r"D:\02_2026년\12_소아비만\발표자료\소아비만_1차후보지역_리스트안_260722.docx"

co = pd.read_csv(os.path.join(OUT, "local_coefs_ob_MGTWR.csv"), dtype={"sgg_cd": str})
enp = pd.read_csv(os.path.join(OUT, "enp_critical_t_ob.csv")).set_index("variable")
mp = pd.read_csv(os.path.join(OUT, "master_panel_2015_2023.csv"), dtype={"sgg_cd": str})
dv = pd.read_csv(os.path.join(OUT, "dependent_vars_2012_2024.csv"), dtype={"sgg_cd": str})
name2code = mp.drop_duplicates("sgg_cd").set_index("sgg_nm")["sgg_cd"].to_dict()
code2name = {v: k for k, v in name2code.items()}

VARS = {"welfare_rate": "수급률", "divorce_rate": "이혼율",
        "youth_net_mig": "청년이동", "apt_ratio": "아파트"}

# 후보지명(docx) -> 229 패널 지역명
MAP = {"옥천군": "옥천군(충북)", "영동군": "영동군(충북)", "아산시": "아산시(충남)",
       "천안시 서북구": "천안시(충남)", "보령시": "보령시(충남)", "서천군": "서천군(충남)",
       "강릉시": "강릉시(강원)", "정선군": "정선군(강원)", "평창군": "평창군(강원)",
       "영월군": "영월군(강원)", "원주시": "원주시(강원)", "강서구": "강서구(서울)",
       "구로구": "구로구(서울)", "노원구": "노원구(서울)", "화성시": "화성시(경기)",
       "남양주시": "남양주시(경기)", "평택시": "평택시(경기)", "인천 서구": "서구(인천)",
       "시흥시": "시흥시(경기)", "인천 남동구": "남동구(인천)", "인천 부평구": "부평구(인천)",
       "부천시 소사구": "부천시(경기)", "목포시": "목포시(전남)", "무안군": "무안군(전남)",
       "울진군": "울진군(경북)", "청송군": "청송군(경북)", "영양군": "영양군(경북)",
       "포항시 북구": "포항시(경북)", "제주시": "제주시(제주)", "서귀포시": "서귀포시(제주)"}


def drivers_for(code):
    d23 = co[(co.sgg_cd == code) & (co.year == 2023)].iloc[0]
    sig23, sus = [], []
    d2023 = co[(co.sgg_cd == code) & (co.year >= 2020)]
    for v, k in VARS.items():
        b, t = d23["b_" + v], d23["t_" + v]
        if abs(t) >= enp.loc[v, "t_crit"]:
            sig23.append(f"{k}{'+' if b > 0 else '−'}({b:.2f})")
        if (d2023["t_" + v].abs() >= enp.loc[v, "t_crit"]).mean() >= 0.5:
            sus.append(f"{k}{'+' if d2023['b_' + v].mean() > 0 else '−'}")
    return ", ".join(sig23) or "—", ", ".join(sus) or "—", len(sus)


def group_of(axis, n):
    if n >= 3:
        return "A (최우선)"
    if n == 2:
        return "A′ (우선)"
    if "규모" in axis and n >= 1:
        return "B (규모 효율)"
    return "C (단일 동인·전략)"

# ---- docx 표 읽기 ----
d = docxlib.Document(SRC_DOCX)
t = d.tables[0]
rows = []
for r in t.rows[1:]:
    c = [x.text.strip().replace("\n", " ") for x in r.cells]
    rows.append(c)  # 연번, 권역, 시도, 시군구, 선정축, 비만율, 과체중, 아동수

recs = []
for c in rows:
    sgg_docx = c[3]
    nm = MAP[sgg_docx]
    code = name2code[nm]
    sig23, sus, n = drivers_for(code)
    recs.append({
        "연번": int(c[0]), "권역": c[1], "시도": c[2], "시군구": sgg_docx,
        "선정 축": c[4], "비만율(%) 2024·EHSA": c[5], "과체중이상율(%) 2024·EHSA": c[6],
        "비만 아동수(명)·EHSA": c[7],
        "MGTWR 유의 동인(2023, β)": sig23,
        "지속 유의 동인(2020–23)": sus,
        "동인 수": n, "우선순위 그룹": group_of(c[4], n)})
df = pd.DataFrame(recs)

# ---- 누락 검토 ----
ob24 = dv[dv.year == 2024].set_index("sgg_cd")["obesity_rate"].to_dict()
cand_names = set(MAP.values())
miss = []
for code in co.sgg_cd.unique():
    nm = code2name.get(code, code)
    if nm in cand_names:
        continue
    _, sus, n = drivers_for(code)
    o = ob24.get(code)
    if o and n >= 2 and o >= 11.3:
        miss.append({"지역": nm, "2024 비만율(%)": round(o, 1),
                     "지속 유의 동인(2020–23)": sus, "동인 수": n})
miss_df = pd.DataFrame(sorted(miss, key=lambda x: -x["2024 비만율(%)"]))

grp = (df.groupby("우선순위 그룹")
         .agg(지역수=("연번", "count"),
              지역=("시군구", lambda s: ", ".join(s)))
         .reset_index())

# ---- 엑셀 ----
xlsx = os.path.join(OUT, "지역선정_MGTWR교차검토_260729.xlsx")
with pd.ExcelWriter(xlsx, engine="openpyxl") as w:
    df.to_excel(w, sheet_name="교차표(30곳)", index=False)
    grp.to_excel(w, sheet_name="우선순위 요약", index=False)
    miss_df.to_excel(w, sheet_name="누락 검토", index=False)
    from openpyxl.styles import Font, PatternFill, Alignment
    for ws in w.book.worksheets:
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="1D3350")
            cell.alignment = Alignment(horizontal="center", vertical="center")
        for col in ws.columns:
            width = max(len(str(c.value or "")) for c in col) + 2
            ws.column_dimensions[col[0].column_letter].width = min(width * 1.7, 46)
        ws.freeze_panes = "A2"
print("saved", xlsx)

# ---- 워드 ----
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docxlib.Document()
st = doc.styles["Normal"]
st.font.name = "맑은 고딕"
st.font.size = Pt(10)
st.element.rPr.rFonts.set(docxlib.oxml.ns.qn("w:eastAsia"), "맑은 고딕")

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run.element.rPr.rFonts.set(docxlib.oxml.ns.qn("w:eastAsia"), "맑은 고딕")
        run.font.color.rgb = RGBColor(0x1D, 0x33, 0x50)
    return p

def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p

def table_from_df(tdf, widths=None, fontsize=8.5):
    t = doc.add_table(rows=1, cols=len(tdf.columns))
    t.style = "Table Grid"
    for i, c in enumerate(tdf.columns):
        cell = t.rows[0].cells[i]
        cell.text = str(c)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(fontsize)
    for _, row in tdf.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(fontsize)
    return t

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("중재프로그램 1차 후보지역(30곳) × MGTWR 환경 동인 교차 검토")
r.bold = True
r.font.size = Pt(16)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("작성일: 2026-07-29 · 작성: 임창민(국립공주대학교 지리학과) · 이현희(건국대학교 지리학과)").font.size = Pt(9)

h("1. 개요와 방법", 1)
para("목적: EHSA 기반 1차 후보지역 30곳(고비율 축 16·규모 축 13·전략 거점 1)에 대해, "
     "다중스케일 시공간 가중회귀(MGTWR)로 추정한 지역별 환경 동인의 유의성을 결합하여 "
     "선정의 타당성을 교차 검증하고 우선순위·개입 방향을 제안한다.")
para("방법: 시군구 패널(229개 × 2015–2023)에 대한 MGTWR 로컬 계수(비만율 모형)에서 "
     "각 후보지의 4개 시기 가변 변수(기초생활수급률·조이혼율·청년순이동률·아파트비율) 유의성을 "
     "ENP 보정 임계 t 기준으로 판정하였다. 2023년 단면과 2020–23년 4개년 중 과반 유의(지속 유의)를 병기한다. "
     "EHSA는 '어디가 부담인가', MGTWR은 '그곳에서 어떤 환경의 힘이 작동하는가'를 답하므로 상호 보완적이다.")

h("2. 핵심 발견 — 선정 축과 환경 동인의 정합", 1)
para("고비율 축 16곳(농촌·산간)은 아파트비율(−) 16곳 전부, 수급률(+) 10곳, 청년이동(−) 6곳이 "
     "지속 유의하여 결핍·지역쇠퇴·주거환경이 실제로 작동하는 지역임이 확인된다. "
     "규모 축 13곳(수도권·대도시)은 13곳 전부 조이혼율(+) 단일 동인으로, 가족구조 경로가 공통이다. "
     "독립적인 두 방법이 같은 공간 구도로 수렴하므로 두 축 구성은 타당하며, "
     "여기에 '중재 가능성(동인)' 축을 더해 우선순위를 정할 수 있다. "
     "전략 거점 원주는 유의 동인이 사실상 없어(지속 기준 청년이동 1개) 통계 근거는 약하다 — "
     "중재 효과 검증을 위한 대조·비교 거점으로 명분화할 것을 권장한다.")

h("3. 우선순위 제안 (부담 × 동인 수)", 1)
table_from_df(grp, fontsize=9)
para("")
para("A(지속 동인 3개): 복합 취약지 — 결핍/가족구조·청년유출·주거환경이 동시에 작동. 최우선 투입.", bold=True)
para("A′(동인 2개): 수급률+아파트(강원·전남·제주) 또는 이혼율+아파트(충남 북부) 조합.")
para("B(규모 축): 동인이 이혼율 하나로 동일 → 표준 프로그램의 대규모 적용. "
     "아동수·EHSA 심화형 기준 화성(603명)·평택(392)·시흥(350) 우선.")
para("C: 아파트(−) 단일 동인 지역(보령·서천·청송·포항 북구)은 주거·활동환경 중심 단일 처방, "
     "원주는 검증 거점 프레임.")

h("4. 동인별 개입 방향", 1)
table_from_df(pd.DataFrame([
    ["수급률(+)", "강원 영동권 4곳, 전남 2곳, 울진·영양, 제주 2곳", "저소득 가정 영양·신체활동 지원(바우처, 급식 질) 상시 프로그램"],
    ["이혼율(+)", "수도권 13곳 전부 + 충청 북부(아산·천안·옥천·영동)", "한부모·맞벌이 가정 아동 돌봄·생활습관 프로그램"],
    ["청년이동(−)", "옥천·영동·영월·울진·청송·영양", "지역활력·정주여건 사업과 아동건강 사업 연계"],
    ["아파트(−)", "고비율 축 전부", "비아파트 주거지 중심 놀이·활동 공간 조성"]],
    columns=["지속 동인", "해당 후보지", "개입 방향"]), fontsize=9)

h("5. 누락 검토 (권역 대표성 규칙 감안)", 1)
para("부담(2024 비만율 전국 평균 이상)과 지속 동인 2개 이상을 갖추었으나 리스트에 없는 지역은 아래와 같다. "
     "대부분 기존 군집의 비대표 지역이나, 진도군(2024 비만율 25.4%, 전국 최고 수준)과 "
     "보은군(전국 유일 동인 4개 전부 지속 유의)은 개별 검토 가치가 있다. "
     "보은군은 같은 충북 남부 군집의 옥천군과 대표 교체 또는 추가를 검토할 만하다.")
table_from_df(miss_df, fontsize=9)

h("6. 부록 — 30개 후보지 전체 교차표", 1)
table_from_df(df.drop(columns=["동인 수"]), fontsize=7.5)

h("7. 유의사항", 1)
para("· 구 단위 후보(천안 서북구·부천 소사구·포항 북구)의 동인은 시 단위 MGTWR 계수로 근사하였다.")
para("· MGTWR 동인은 연관이며 인과가 아니다 — '해당 지역에서 해당 환경 격차가 비만 격차와 함께 움직인다'로 해석한다.")
para("· 수치 재현: output/local_coefs_ob_MGTWR.csv, enp_critical_t_ob.csv, scripts/10_region_report.py")

wpath = os.path.join(OUT, "지역선정_MGTWR교차검토_260729.docx")
doc.save(wpath)
print("saved", wpath)
